import os
import io

import streamlit as st
from docx import Document
from PIL import Image

# Titolo della scheda del browser e layout della pagina.
# layout="centered" = colonna centrale stretta, "wide" = tutta la larghezza.
st.set_page_config(page_title="Tesina", layout="centered")

# Percorso alla cartella con tutti i file .docx e l'immagine di copertina.
# Se sposti la cartella DEFINITIVA, modifica solo questa riga.
BASE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "DEFINITIVA")

# --- ELENCO DEI CAPITOLI ---
# Ogni voce: (nome_file, numero_romano)
#   nome_file    → usato per costruire il percorso dei .docx
#                  es. "terzo" → "capitolo terzo testo.docx" / "CAPITOLO TERZO.docx"
#   numero_romano → etichetta mostrata nella sidebar e come titolo
#
# Per AGGIUNGERE un capitolo: aggiungi una riga e crea i due file corrispondenti:
#   DEFINITIVA/capitolo tredicesimo testo.docx   ← solo testo
#   DEFINITIVA/CAPITOLO TREDICESIMO.docx         ← versione con immagini
#
# Per RIMUOVERE un capitolo: cancella la riga.
# Per RIORDINARE: cambia l'ordine delle righe.
CHAPTERS = [
    ("primo",      "I"),
    ("secondo",    "II"),
    ("terzo",      "III"),
    ("quarto",     "IV"),
    ("quinto",     "V"),
    ("sesto",      "VI"),
    ("settimo",    "VII"),
    ("ottavo",     "VIII"),
    ("nono",       "IX"),
    ("decimo",     "X"),
    ("undicesimo", "XI"),
    ("dodicesimo", "XII"),
]


# --- LETTURA TESTO DA .docx ---
# Legge tutti i paragrafi e li converte in Markdown.
# Se c'è un errore nel file (es. docx corrotto o formato insolito),
# mostra un messaggio in corsivo invece di crashare l'app.
# Il risultato è in cache: il file viene riletto solo al riavvio dell'app.
@st.cache_data
def read_text(filename: str, skip_chapter_label: bool = False) -> str:
    # skip_chapter_label=True attiva due filtri usati solo per i capitoli:
    #   1. Salta righe tipo "CAPITOLO PRIMO", "CAPITOLO SESTO" (ridondanti con st.title)
    #   2. Nelle prime 10 righe, salta righe in maiuscolo già contenute
    #      in una riga precedente (gestisce titoli duplicati nei docx)
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        return f"_File non trovato: {filename}_"
    try:
        doc = Document(path)
        lines = []
        seen_caps = []  # traccia heading già viste per deduplicare

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            if skip_chapter_label:
                # Salta "CAPITOLO [nome]" — es. "CAPITOLO PRIMO", "CAPITOLO SESTO"
                words = text.split()
                if words[0].upper() == "CAPITOLO" and len(words) <= 3:
                    continue

                # Nelle prime 10 righe, deduplica heading in maiuscolo:
                # salta se il testo è già interamente contenuto in una riga vista prima
                if len(lines) < 10 and text.upper() == text:
                    norm = text.upper()
                    if any(norm in seen for seen in seen_caps):
                        continue
                    seen_caps.append(norm)

            # Gestione stili di titolo (Heading). Se lo stile non è accessibile
            # si usa testo normale come fallback.
            try:
                style = p.style.name
            except Exception:
                style = ""
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading"):
                lines.append(f"### {text}")
            else:
                lines.append(text)

        if not lines:
            return "_Il file non contiene testo leggibile (potrebbe usare tabelle o caselle di testo non supportate)._"
        return "\n\n".join(lines)
    except Exception as e:
        return f"_Errore nella lettura del file `{filename}`: {e}_"


# --- ESTRAZIONE IMMAGINI DA .docx ---
# Estrae tutte le immagini incorporate nel file .docx.
# Le restituisce come lista di oggetti PIL.Image pronti per st.image().
# Errori su singole immagini vengono saltati silenziosamente.
@st.cache_data
def extract_images(filename: str) -> list:
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        return []
    try:
        doc = Document(path)
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    blob = rel.target_part.blob
                    images.append(Image.open(io.BytesIO(blob)))
                except Exception:
                    pass
        return images
    except Exception:
        return []


# --- MENU DI NAVIGAZIONE ---
# L'ordine delle voci rispecchia l'ordine di questa lista.
# Per aggiungere una sezione fissa (es. "Appendice"), aggiungi la stringa
# nella posizione desiderata e aggiungi il blocco elif corrispondente in fondo.
nav_items = ["Copertina", "Dedica", "Indice", "Introduzione"]
for _, roman in CHAPTERS:
    nav_items.append(f"Capitolo {roman}")
nav_items += ["Conclusioni", "Ringraziamenti"]

with st.sidebar:
    st.markdown("## Navigazione")
    selection = st.radio("", nav_items, label_visibility="collapsed")


# =============================================================================
# RENDERING DELLE PAGINE
#
# Struttura standard di ogni sezione:
#   1. st.title(...)      → titolo della pagina
#   2. st.markdown(...)   → testo estratto dal .docx di solo testo
#   3. st.divider()       → separatore visivo (solo se ci sono immagini)
#   4. st.image(...)      → immagini estratte dal .docx con foto
#
# Per aggiungere una sezione fissa nuova (es. "Appendice"):
#   1. Aggiungi "Appendice" alla lista nav_items sopra
#   2. Crea DEFINITIVA/appendice.docx
#   3. Aggiungi qui sotto:
#      elif selection == "Appendice":
#          st.title("Appendice")
#          st.markdown(read_text("appendice.docx"))
# =============================================================================

if selection == "Copertina":
    # Immagine PNG di copertina + eventuale testo dal docx.
    # Per cambiare immagine: sostituisci "le_mille_e_una_notte.png" in DEFINITIVA.
    cover_path = os.path.join(BASE, "le_mille_e_una_notte.png")
    if os.path.exists(cover_path):
        st.image(cover_path, use_container_width=True)
    text = read_text("copertina con simbolo.docx")
    if text:
        st.markdown(text)

elif selection == "Dedica":
    st.title("Dedica")
    # File: DEFINITIVA/dedica.docx
    st.markdown(read_text("dedica.docx"))

elif selection == "Indice":
    st.title("Indice")
    # File: DEFINITIVA/INDICE.docx
    st.markdown(read_text("INDICE.docx"))

elif selection == "Introduzione":
    st.title("Introduzione")
    # Testo da:    DEFINITIVA/Introduzione testo.docx
    # Immagini da: DEFINITIVA/INTRODUZIONE.docx
    st.markdown(read_text("Introduzione testo.docx"))
    images = extract_images("INTRODUZIONE.docx")
    if images:
        st.divider()
        for img in images:
            st.image(img, use_container_width=True)

elif selection.startswith("Capitolo"):
    # Ricava il nome del file dal numero romano selezionato.
    # Es. "Capitolo III" → roman="III" → name="terzo"
    #   testo:    DEFINITIVA/capitolo terzo testo.docx
    #   immagini: DEFINITIVA/CAPITOLO TERZO.docx
    roman = selection.split(" ", 1)[1]
    name = next(n for n, r in CHAPTERS if r == roman)
    st.title(selection)
    st.markdown(read_text(f"capitolo {name} testo.docx", skip_chapter_label=True))
    images = extract_images(f"CAPITOLO {name.upper()}.docx")
    if images:
        st.divider()
        for img in images:
            st.image(img, use_container_width=True)

elif selection == "Conclusioni":
    st.title("Conclusioni")
    # Testo da:    DEFINITIVA/CONCLUSIONI.docx
    # Immagini da: DEFINITIVA/CONCLUSIONI FOTO.docx
    st.markdown(read_text("CONCLUSIONI.docx"))
    images = extract_images("CONCLUSIONI FOTO.docx")
    if images:
        st.divider()
        for img in images:
            st.image(img, use_container_width=True)

elif selection == "Ringraziamenti":
    st.title("Ringraziamenti")
    # File: DEFINITIVA/ringraziamenti finali.docx
    st.markdown(read_text("ringraziamenti finali.docx"))
